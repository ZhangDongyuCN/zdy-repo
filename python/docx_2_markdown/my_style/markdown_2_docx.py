# -*- coding: utf-8 -*-

import os
import re
import glob
import yaml
import shutil
import docx
from docx.shared import Pt, RGBColor
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def write_1eading_1(document, line):
    line = line.replace('# ', '').rstrip('\n')
    paragraph = document.add_paragraph('')
    run = paragraph.add_run(line)
    run.font.name = '微软雅黑'
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0x70, 0x30, 0xA0)


def write_1eading_2(document, line):
    line = line.replace('## ', '').rstrip('\n')
    paragraph = document.add_paragraph('')
    run = paragraph.add_run(line)
    run.font.name = '微软雅黑'
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)


def write_1eading_3(document, line):
    line = line.replace('### ', '').rstrip('\n')
    paragraph = document.add_paragraph('')
    run = paragraph.add_run(line)
    run.font.name = '微软雅黑'
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0xB0, 0x50)


def write_1eading_4(document, line):
    line = line.replace('#### ', '').rstrip('\n')
    paragraph = document.add_paragraph('')
    run = paragraph.add_run(line)
    run.font.name = '微软雅黑'
    run.font.size = Pt(10.5)
    run.bold = True
    run.font.color.rgb = RGBColor(0xC5, 0x5A, 0x11)


def set_cell_background_color(cell, color_str):
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value=color_str))
    cell._tc.get_or_add_tcPr().append(shading_elm_1)


def write_block_code(f_md, document):
    block_code = ''
    line = f_md.readline()
    while line.lstrip(' ').find('```') != 0:
        block_code += line
        line = f_md.readline()
    block_code = block_code.strip('\n')

    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.cell(0, 0).text = block_code
    set_cell_background_color(table.cell(0, 0), 'F2F2F2')


def write_list_bullet(document, line, md_path):
    line = line.replace('- ', '')
    write_normal_text(document, line, md_path, 'List Bullet')


def write_list_number(document, line, md_path):
    line = re.sub(r'\d+. ', '', line)
    write_normal_text(document, line, md_path, 'List Number')


def write_line_code(paragraph, line, i, start_idx):
    # 写入普通文本
    normal_text = line[start_idx:i]
    paragraph.add_run(normal_text)

    # 写入行代码
    line_code = ''
    i += 1
    while line[i] != '`':
        line_code += line[i]
        i += 1
    run = paragraph.add_run(line_code)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    start_idx = i + 1

    return i, start_idx


def write_bold_italic_text(paragraph, line, i, start_idx):
    old_i = i

    # 如果是加粗
    if line[i + 1] == '*':
        # 写入普通文本
        normal_text = line[start_idx:i]
        paragraph.add_run(normal_text)

        # 写入加粗字体
        bold_code = ''
        i += 2
        while line[i] != '*':
            bold_code += line[i]
            i += 1
            if i == len(line):
                return old_i, start_idx
        i += 1
        run = paragraph.add_run(bold_code)
        run.bold = True
        start_idx = i + 1
    # 如果是倾斜
    else:
        # 写入普通文本
        normal_text = line[start_idx:i]
        paragraph.add_run(normal_text)

        # 写入加粗字体
        bold_code = ''
        i += 1
        while line[i] != '*':
            bold_code += line[i]
            i += 1
            if i == len(line):
                return old_i, start_idx
        run = paragraph.add_run(bold_code)
        run.italic = True
        start_idx = i + 1

    return i, start_idx


def write_image(document, paragraph, line, i, start_idx, md_path):
    # 如果是图片
    if i < len(line) - 2 and line[i + 1] == '[':
        # 查找图片的说明
        end_pos = line.find(']', i + 2)

        # + 排除：“xxxx![xxxx[xxxx]xxx”
        temp_pos = line.find('[', i + 2)
        if (end_pos == -1) or (temp_pos != -1 and end_pos > temp_pos):
            return i, start_idx

        # 查找图片路径（仅支持本地图片）
        # + 排除：“xxxx![xxxx]”或“xxxx![xxxx]xxx”
        start_pos = end_pos + 1
        if (start_pos == len(line)) or (line[start_pos] != '('):
            return i, start_idx
        # + url
        end_pos = line.find(')', start_pos + 1)
        if end_pos == -1:  # 如果没找到，就不能判定为url
            return i, start_idx
        img_path = line[start_pos + 1:end_pos]
        if not os.path.exists(img_path):
            img_path = os.path.join(os.path.dirname(md_path), img_path)

        # 写入普通文本
        normal_text = line[start_idx:i]
        paragraph.add_run(normal_text)

        # 添加图片
        document.add_picture(img_path)

        i = end_pos
        start_idx = end_pos + 1

        return i, start_idx
    # 如果是普通感叹号
    else:
        return i, start_idx


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def write_hyperlink(paragraph, line, i, start_idx):
    # 查找超链接的名字
    end_pos = line.find(']', i + 1)

    # + 排除：“xxxx[xxxx[xxxx]xxx”
    temp_pos = line.find('[', i + 1)
    if (end_pos == -1) or (temp_pos != -1 and end_pos > temp_pos):
        return i, start_idx
    text = line[i + 1:end_pos]

    # 查找超链接的链接
    # + 排除：“xxxx[xxxx]”或“xxxx[xxxx]xxx”
    start_pos = end_pos + 1
    if (start_pos == len(line)) or (line[start_pos] != '('):
        return i, start_idx
    # + url
    end_pos = line.find(')', start_pos + 1)
    if end_pos == -1:  # 如果没找到，就不能判定为超链接
        return i, start_idx
        # raise ValueError('you markdown format may be not right, current line is: {}'.format(line))
    url = line[start_pos + 1:end_pos]

    # 写入普通文本
    normal_text = line[start_idx:i]
    paragraph.add_run(normal_text)

    # 添加超链接
    add_hyperlink(paragraph, text, url)

    i = end_pos
    start_idx = end_pos + 1

    return i, start_idx


def write_normal_text(document, line, md_path, style=None):
    paragraph = document.add_paragraph('', style)

    i = 0
    line_len = len(line)
    start_idx = 0
    while i < line_len:
        if line[i] == '`':
            i, start_idx = write_line_code(paragraph, line, i, start_idx)
        elif line[i] == '*':
            i, start_idx = write_bold_italic_text(paragraph, line, i, start_idx)
        elif line[i] == '!':
            i, start_idx = write_image(document, paragraph, line, i, start_idx, md_path)
        elif line[i] == '[':
            i, start_idx = write_hyperlink(paragraph, line, i, start_idx)

        i += 1

    # 写入最后的normal文本
    paragraph.add_run(line[start_idx:])


def markdown_2_docx(md_path, docx_save_path):
    # 创建文档，并设置基本样式
    document = docx.Document()
    document.styles['Normal'].font.name = '微软雅黑'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    document.styles['Normal'].font.size = Pt(10.5)
    document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

    try:
        # 逐行解析markdown，生成word样式
        with open(md_path, 'rt', encoding='utf-8') as f_md:
            line = f_md.readline()
            while line:
                line = line.rstrip('\n')

                # 判断是不是空
                if line.isspace():
                    document.add_paragraph('')

                # 一级标题
                elif line.find('# ') == 0:
                    write_1eading_1(document, line)

                # 二级标题
                elif line.find('## ') == 0:
                    write_1eading_2(document, line)

                # 三级标题
                elif line.find('### ') == 0:
                    write_1eading_3(document, line)

                # 四级标题
                elif line.find('#### ') == 0:
                    write_1eading_4(document, line)

                # 项目符号
                elif line.lstrip(' ').find('- ') == 0:
                    write_list_bullet(document, line, md_path)

                # 列表
                elif re.match(r'\d+. ', line.lstrip(' ')) is not None:
                    write_list_number(document, line, md_path)

                # 块代码
                elif line.lstrip(' ').find('```') == 0:
                    write_block_code(f_md, document)

                # 正文文本
                else:
                    write_normal_text(document, line, md_path)

                # 读取下一行
                line = f_md.readline()

            document.save(docx_save_path)
    except Exception as e:
        print('[ERROR]', e)
        print('[ERROR] error in markdown_2_docx(), line is: {}'.format(line))


def enter_fun():
    # 解析配置文件
    config = yaml.load(open('./docx_2_markdown_config.yaml', 'rt', encoding='utf-8'), Loader=yaml.FullLoader)
    word_mode = config["word_mode"]
    src_path = config["src_path"]
    save_path = config["save_path"]

    if word_mode == 1:
        markdown_2_docx(src_path, save_path)
    if word_mode == 2:
        # 创建目录
        if os.path.exists(save_path):
            shutil.rmtree(save_path)
        os.makedirs(save_path)

        # 列出所有markdown文档
        md_files = glob.glob(os.path.join(src_path, '*.md'))

        # markdown 2 docx
        for md_file in md_files:
            docx_file = os.path.join(save_path, os.path.splitext(os.path.basename(md_file))[0] + '.docx')
            markdown_2_docx(md_file, docx_file)
